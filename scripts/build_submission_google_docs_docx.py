#!/usr/bin/env python3
"""Build an editable, Google Docs-targeted DOCX from the current manuscript PDF.

The Prism PDF is the editorial source of truth. Text is recovered by position,
margin line numbers and figure placeholders are removed, and the frozen main
figure assets are inserted as accessible images. Equations and the final test
status are normalized explicitly where PDF text extraction is ambiguous.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re

import pdfplumber
from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/manuscripts/fire_vase_developmental_morphology/main-22.pdf"
FIGURE_DIR = ROOT / "figures/submission"
DRAFT = Path("/private/tmp/fire_vase_manuscript_google_docs_draft.docx")

PLACEHOLDER_PREFIXES = (
    "Validated v2 Figure",
    "Add Figure_",
    "Open the validated repository figure",
)

DEHYPHENATE = {
    "developmental", "geometry", "earlier", "wildfire", "visualization",
    "reviewer", "nonnumeric", "remote-sensing",
}

H1_EXACT = {
    "One-Sentence Summary", "Abstract", "1 Introduction", "2 Results",
    "3 Discussion", "4 Materials and Methods", "References and Notes",
    "References", "Acknowledgments", "Supplementary Materials",
}


@dataclass
class Line:
    page: int
    top: float
    x0: float
    size: float
    text: str


@dataclass
class Block:
    kind: str
    text: str
    page: int
    top: float


def join_text(left: str, right: str) -> str:
    left = left.rstrip()
    right = right.lstrip()
    if left.endswith("-") and right:
        left_word = re.search(r"([A-Za-z]+)-$", left)
        right_word = re.match(r"([A-Za-z]+)", right)
        if left_word and right_word:
            combined = left_word.group(1) + right_word.group(1)
            if combined.lower() in DEHYPHENATE:
                return left[:-1] + right
        return left + right
    return f"{left} {right}".strip()


def clean_text(text: str) -> str:
    replacements = {
        "1Environmental": "¹ Environmental",
        "Ph.D.1": "Ph.D.¹",
        "Ph.D. 1": "Ph.D.¹",
        "pthe same records": "the same records",
        "fast fires— 2 defined": "fast fires—defined",
        "16. 2 km in one day": "16.2 km² in one day",
        "A(T )": "A(T)",
        "km2": "km²",
        "R2": "R²",
        "R 2": "R²",
        "Day- t": "Day-t",
        "day- t": "day-t",
        "km 2": "km²",
        "day − 1": "day⁻¹",
        "2.7%": "2.7%",
        "proportional to A(t)/A(T)": "proportional to √[A(t) / A(T)]",
        "ring radius (cid:112) is proportional to A (t ) /A (T )":
            "ring radius is proportional to √[A(t) / A(T)]",
        "it is below 2 2 1% among fires no larger than 1 km, rises above 18% among fires of 10–100 km,":
            "it is below 1% among fires no larger than 1 km², rises above 18% among fires of 10–100 km²,",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"(?<=\d)\.\s+(?=\d)", ".", text)
    text = re.sub(r"\(\s+", "(", text)
    text = text.replace("= − ", "= −")
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def extract_lines() -> list[Line]:
    output: list[Line] = []
    with pdfplumber.open(SOURCE) as pdf:
        for page_number, page in enumerate(pdf.pages, 1):
            words = page.extract_words(x_tolerance=1, extra_attrs=["fontname", "size"])
            groups: list[list[dict]] = []
            for word in sorted(words, key=lambda item: (item["top"], item["x0"])):
                if groups and abs(word["top"] - groups[-1][0]["top"]) <= 2.0:
                    groups[-1].append(word)
                else:
                    groups.append([word])
            for group in groups:
                group.sort(key=lambda item: item["x0"])
                text = " ".join(word["text"] for word in group).strip()
                x0 = min(word["x0"] for word in group)
                top = min(word["top"] for word in group)
                if not text:
                    continue
                if text.startswith(PLACEHOLDER_PREFIXES):
                    continue
                if re.fullmatch(r"\d+", text) and (
                    x0 < 65 or x0 > 545 or top > 730 or (page_number == 1 and top < 200)
                ):
                    continue
                if page_number == 1 and text == "Ty Tuff, Ph.D.":
                    text += "¹"
                if page_number == 1 and text == "1" and 200 < top < 250:
                    text = "¹"
                if page_number == 1 and text.startswith("Environmental Data Science"):
                    text = "¹ " + text
                if page_number == 1 and text.startswith("1 Environmental Data Science"):
                    text = "¹" + text[1:]
                size = max((word.get("size", 0) for word in group), default=0)
                output.append(Line(page_number, top, x0, size, clean_text(text)))
    return output


def is_h2(text: str) -> bool:
    return bool(re.match(r"^[234]\.\d+\s+\S", text))


def classify(line: Line) -> str:
    if line.page == 1 and line.size >= 17:
        return "title"
    if line.page == 1 and line.top < 330:
        return "meta"
    if line.text in H1_EXACT:
        return "h1"
    if is_h2(line.text):
        return "h2"
    if re.match(r"^Figure [1-5]:", line.text):
        return "caption"
    if re.match(r"^\[\d+\]", line.text):
        return "reference"
    return "body"


def blocks_from_lines(lines: list[Line]) -> list[Block]:
    blocks: list[Block] = []
    reference_mode = False
    caption_mode = False
    for line in lines:
        kind = classify(line)
        if kind in {"h1", "h2", "title"}:
            reference_mode = False
            caption_mode = False
        elif kind == "reference":
            reference_mode = True
            caption_mode = False
        elif kind == "caption":
            caption_mode = True
        elif reference_mode and kind == "body":
            kind = "reference_continuation"
        elif caption_mode and kind == "body":
            kind = "caption_continuation"

        if kind == "reference_continuation" and blocks and blocks[-1].kind == "reference":
            blocks[-1].text = join_text(blocks[-1].text, line.text)
            continue
        if kind == "caption_continuation" and blocks and blocks[-1].kind == "caption":
            blocks[-1].text = join_text(blocks[-1].text, line.text)
            continue

        if blocks:
            previous = blocks[-1]
            same_page_gap = line.top - previous.top if line.page == previous.page else None
            if kind == "body" and previous.kind == "h2" and line.size >= 11.9 and same_page_gap is not None and same_page_gap < 28:
                previous.text = join_text(previous.text, line.text)
                previous.top = line.top
                continue
            if (
                kind == "meta" and previous.kind == "meta"
                and same_page_gap is not None and same_page_gap < 28
                and not line.text.startswith(("ORCID:", "Correspondence:"))
                and not previous.text.startswith(("ORCID:", "Correspondence:"))
            ):
                previous.text = join_text(previous.text, line.text)
                previous.top = line.top
                continue
            if kind == "h2" and previous.kind == "h2" and same_page_gap is not None and same_page_gap < 28:
                previous.text = join_text(previous.text, line.text)
                previous.top = line.top
                continue
            if kind in {"body", "caption"} and previous.kind == kind:
                if same_page_gap is not None and same_page_gap < 27:
                    previous.text = join_text(previous.text, line.text)
                    previous.top = line.top
                    continue
                if line.page == previous.page + 1 and not re.search(r"[.!?]$", previous.text):
                    previous.text = join_text(previous.text, line.text)
                    previous.page = line.page
                    previous.top = line.top
                    continue
        if kind in {"reference_continuation", "caption_continuation"}:
            kind = "body"
        blocks.append(Block(kind, line.text, line.page, line.top))
    return blocks


SECTION_42 = [
    "For nonnegative daily increments gⱼ, reconstructed area was S = ∑ⱼ gⱼ and normalized growth allocation was pⱼ = gⱼ/S. For a positive-area fire, normalized VASE width at developmental time t was",
    "w(t) = √[A(t) / A(T)].    (1)",
    "Here A(t) is cumulative reconstructed area and A(T) = S. Every VASE therefore terminates at width 1. Width is proportional, up to a constant, to the radius of an equal-area disk; it is not a geographic silhouette.",
]

SECTION_43 = [
    "In the primary consecutive cohort, each daily increment occupied one equal-width interval of relative developmental time. We interpolated cumulative allocation at 21 fixed boundaries and calculated the difference between adjacent boundaries, producing 20 nonnegative growth shares that sum to one. Interpolation standardizes the representation but does not create observations. The displayed VASE uses cumulative allocation to determine radius, whereas these 20 incremental allocation bins were the only inputs to the primary PCA [10]. Final area, duration, observation count, true observed daily peak, mean growth, slenderness, scalar developmental traits, and weather were excluded. Each bin was centered and scaled using its training-set mean and standard deviation; near-constant columns received a scale factor of one. Because the shares form a composition, the resulting PCA is a standardized Euclidean coordinate representation rather than a unique geometry for the histories. Five axes were retained.",
    "For the compositional sensitivity, allocations were transformed as hᵢⱼ = √pᵢⱼ. Each transformed row therefore had squared Euclidean norm one. Columns were mean centered but not variance scaled, preserving Hellinger distances up to a constant. PCA was then refitted to the same 10,246 fires. Components were matched by maximum absolute score correlation and oriented to positive aligned correlation. Global and local correspondence used the same Procrustes, score-subspace, pairwise-distance, nearest-neighbor, and exemplar-tail diagnostics as the observation-depth analysis.",
    "Shape-oriented scalar traits were calculated after coordinate fitting. Front loading was allocation in the first half of relative development; late allocation was growth in the final quarter. Observed daily peak was the largest dated increment. Shannon entropy was −∑ⱼ pⱼ log pⱼ, with zero-probability terms omitted, and normalized entropy divided by log n for n > 1 [23]. Normalized first and second differences described changes between successive allocation bins and were not interpreted as physical velocity or acceleration.",
    "Interpretive landmarks were assigned by declared rules after excluding invalid, one-observation, two-observation, and gappy histories. Multiple detected pulses required at least two local maxima with prominence at least 20% of the observed peak. Reactivation required growth to return above 25% of peak after two consecutive subthreshold observed days. Late peak, front-loaded taper, and distributed growth used prespecified timing and allocation thresholds. These labels were not PCA inputs or statistically inferred natural classes.",
]

SECTION_46 = [
    "State models used only transitions separated by exactly one calendar day and required an observation on the previous day for every comparison. The common cohort contained 87,944 transitions from 31,700 fires. The response was log[1 + gₜ₊₁ (km²)]. The autoregressive state baseline contained current growth, previous-day growth, cumulative reconstructed area, and elapsed time. Models then added current-day weather and interactions between weather and state. All models used the same observations, training-only scaling, the prespecified ridge penalty, and the same validation blocks as the event-level models.",
    "VPD-specific validation removed or added the interactions of VPD with current and previous growth while retaining the other interaction terms. Paired bootstrap intervals tested changes in held-out performance using fixed predictions. Coefficients expressed in the original units used standard errors clustered by fire and remained conditional on the fitted design and penalty. Refits by region, year, season, fire size, and training fold assessed heterogeneity. Tables of VPD by fire state reported both transitions and distinct fires without extrapolating a response surface beyond the observed data.",
]

SECTION_47 = [
    "Matching used either core event-mean weather variables or the 20 morphology bins, standardized within the 9,212-fire cohort. Potential partners were the 10 nearest neighbors within the same region, season, duration, and observation-count group. Pairs with an area ratio greater than 2 or a root-mean-square standardized distance greater than 0.5 were rejected. The remaining pairs were sorted by distance and identifiers and selected in order without reusing a fire; differences in the other representation did not influence selection. Sensitivity analyses varied the number of candidate neighbors, distance measure, and maximum acceptable distance. Conditional reference distributions shuffled the other representation within the same adjustment groups and an additional area class, ⌊log₂(area)⌋.",
]

SECTION_48 = [
    "All analyses use the analysis-ready v0.1 FIRED/gridMET data lake, with no synthetic fallback. The validated baseline SHA is 18c923cb0c82bf9f66567b62a3491ac30a28c369; the submission freeze manifest records subsequent code and output hashes. Configuration config/analysis_v2.json uses seed 20260828. Checks of scientific invariants, reference-data conservation, manuscript claims, publication hashes, deterministic regeneration, and figure copies passed. The full repository suite passes: 152 passed, 2 skipped, and 125 warnings, with no collection errors or excluded modules.",
    "Daily gridMET and satellite burn dates have timing and spatial limitations. Nearest-cell exposures omit within-fire heterogeneity, directional wind, fuel continuity, terrain, suppression, and active-edge conditions. The primary and weather populations are selected by observation support and completeness. All model and matching results are observational and retrospective.",
]


def replace_sections(blocks: list[Block]) -> list[Block]:
    output: list[Block] = []
    index = 0
    overrides = {
        "4.2": SECTION_42,
        "4.3": SECTION_43,
        "4.6": SECTION_46,
        "4.7": SECTION_47,
        "4.8": SECTION_48,
    }
    while index < len(blocks):
        block = blocks[index]
        match = re.match(r"^(4\.[23678])\b", block.text) if block.kind == "h2" else None
        if not match:
            output.append(block)
            index += 1
            continue
        section = match.group(1)
        output.append(block)
        for text in overrides[section]:
            output.append(Block("equation" if text.startswith("w(t)") else "body", text, block.page, block.top))
        index += 1
        while index < len(blocks):
            candidate = blocks[index]
            if candidate.kind in {"h1", "h2"}:
                break
            index += 1
    return output


def set_font(run, name: str = "Arial") -> None:
    run.font.name = name
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    fonts = run._element.rPr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run._element.rPr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after, color in [
        ("Heading 1", 20, 20, 6, RGBColor(0, 0, 0)),
        ("Heading 2", 16, 18, 6, RGBColor(0, 0, 0)),
        ("Heading 3", 14, 16, 4, RGBColor(67, 67, 67)),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = "Arial"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.bold = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.line_spacing = 1.15
    caption.paragraph_format.keep_together = True

    reference = doc.styles.add_style("Reference", 1)
    reference.base_style = normal
    reference.font.name = "Arial"
    reference._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    reference._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    reference._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    reference.font.size = Pt(10)
    reference.paragraph_format.left_indent = Inches(0.3)
    reference.paragraph_format.first_line_indent = Inches(-0.3)
    reference.paragraph_format.space_after = Pt(7)
    reference.paragraph_format.line_spacing = 1.15


def set_image_alt_text(paragraph, title: str, description: str) -> None:
    doc_pr = paragraph._p.xpath(".//wp:docPr")
    if doc_pr:
        doc_pr[0].set("title", title)
        doc_pr[0].set("descr", description)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_font(run)
    run.font.size = Pt(26)
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_figure(doc: Document, number: int) -> None:
    path = FIGURE_DIR / f"Figure_{number}.png"
    with Image.open(path) as image:
        width_px, height_px = image.size
    width = 6.5
    height = width * height_px / width_px
    if height > 6.2:
        height = 6.2
        width = height * width_px / height_px
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width), height=Inches(height))
    set_image_alt_text(
        paragraph,
        f"Figure {number}",
        f"Fire VASE manuscript Figure {number}; full description follows in the editable caption.",
    )


def add_body(doc: Document, text: str, style: str = "Normal") -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.add_run(text)


def build() -> None:
    blocks = replace_sections(blocks_from_lines(extract_lines()))
    doc = Document()
    configure_document(doc)
    current_figure = 0
    for block in blocks:
        text = clean_text(block.text)
        if block.kind == "title":
            add_title(doc, text)
        elif block.kind == "meta":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(text)
            set_font(run)
            run.font.size = Pt(11)
        elif block.kind == "h1":
            doc.add_heading(text, level=1)
        elif block.kind == "h2":
            doc.add_heading(text, level=2)
        elif block.kind == "caption":
            match = re.match(r"^Figure ([1-5]):\s*(.*)", text)
            if match:
                current_figure = int(match.group(1))
                add_figure(doc, current_figure)
                paragraph = doc.add_paragraph(style="Caption")
                label = paragraph.add_run(f"Figure {current_figure}: ")
                label.bold = True
                body = paragraph.add_run(match.group(2))
                body.bold = False
            else:
                add_body(doc, text, "Caption")
        elif block.kind == "reference":
            add_body(doc, text, "Reference")
        elif block.kind == "equation":
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(text)
            set_font(run)
            run.font.size = Pt(11)
        else:
            add_body(doc, text)

    if current_figure != 5:
        raise RuntimeError(f"Expected five main figures, inserted {current_figure}")
    DRAFT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DRAFT)
    print(DRAFT)


if __name__ == "__main__":
    build()
